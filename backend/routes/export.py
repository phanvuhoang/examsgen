import io
import json
from datetime import datetime
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.database import get_db
from backend.models import ExportRequest

router = APIRouter(prefix="/api/export", tags=["export"])


def _add_mcq_to_doc(doc, content):
    questions = content.get("questions", [])
    for q in questions:
        p = doc.add_paragraph()
        run = p.add_run(f"Question {q.get('number', '')}")
        run.bold = True
        run.font.size = Pt(11)

        if q.get("scenario"):
            doc.add_paragraph(q["scenario"])
        doc.add_paragraph(q.get("question", ""))

        for letter in ["A", "B", "C", "D"]:
            opt = q.get("options", {}).get(letter, {})
            prefix = "[KEY] " if opt.get("is_key") else ""
            doc.add_paragraph(f"  {letter}. {prefix}{opt.get('text', '')}")

        # Answer explanations
        p = doc.add_paragraph()
        run = p.add_run("Answer Explanations:")
        run.bold = True
        run.font.size = Pt(10)
        for letter in ["A", "B", "C", "D"]:
            opt = q.get("options", {}).get(letter, {})
            key_marker = " ✓" if opt.get("is_key") else ""
            doc.add_paragraph(
                f"  {letter}{key_marker}: {opt.get('explanation', '')}",
                style="List Bullet",
            )

        if q.get("regulation_refs"):
            p = doc.add_paragraph()
            run = p.add_run("Regulation References: ")
            run.italic = True
            p.add_run(", ".join(q["regulation_refs"]))

        doc.add_paragraph("")  # spacer


def _add_scenario_to_doc(doc, content):
    p = doc.add_paragraph()
    run = p.add_run(f"Question {content.get('question_number', '')} — {content.get('sac_thue', '')} ({content.get('marks', '')} marks)")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(content.get("scenario", ""))

    for sq in content.get("sub_questions", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{sq.get('label', '')} ({sq.get('marks', '')} marks)")
        run.bold = True
        doc.add_paragraph(sq.get("question", ""))

    # Marking scheme
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Marking Scheme")
    run.bold = True
    run.font.size = Pt(11)

    for sq in content.get("sub_questions", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{sq.get('label', '')}")
        run.bold = True
        if sq.get("answer"):
            doc.add_paragraph(sq["answer"])
        for ms in sq.get("marking_scheme", []):
            doc.add_paragraph(
                f"• {ms.get('point', '')} [{ms.get('mark', '')} mark(s)]",
            )

    doc.add_paragraph("")


@router.post("/word")
def export_to_word(req: ExportRequest):
    if not req.question_ids:
        raise HTTPException(status_code=400, detail="No question IDs provided")

    with get_db() as conn:
        cur = conn.cursor()
        placeholders = ",".join(["%s"] * len(req.question_ids))
        cur.execute(
            f"SELECT id, question_type, content_json FROM questions WHERE id IN ({placeholders}) "
            f"ORDER BY question_part, question_number",
            req.question_ids,
        )
        rows = cur.fetchall()

    if not rows:
        raise HTTPException(status_code=404, detail="No questions found")

    doc = Document()

    # Style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Header
    h = doc.add_heading("ACCA TX(VNM) — Generated Exam Questions", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for row_id, q_type, content in rows:
        if isinstance(content, str):
            content = json.loads(content)
        if q_type == "MCQ":
            _add_mcq_to_doc(doc, content)
        elif q_type in ("SCENARIO_10", "LONGFORM_15"):
            _add_scenario_to_doc(doc, content)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=ExamsGen_{datetime.now().strftime('%Y%m%d')}.docx"},
    )
